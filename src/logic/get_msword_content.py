from docx import Document

data={}
wordDoc = Document('files/Semesterapparat - Anmeldung.docx')
paragraphs = wordDoc.tables
for table in paragraphs:
    for column in table.columns:
        cellcount=0
        for cell in column.cells:
            if cellcount<12:
                cellcount+=1
            print(f'cell:{cell.text}')
        
        
    # print(f'paragraphs[{i}]: {paragraphs[i]}')
    # data[i] = paragraphs[i]
    
# for i in range(0, len(paragraphs)):
# for i in range(2, len(paragraphs)):
#     data[i] = paragraphs[i]
    
print(data)

# for table in wordDoc.tables:
#     for row in table.rows:
#         print('---')
#         for cell in row.cells:
#             print(f'cell:{cell.text}')
            


